import sys
from datetime import datetime
import pytz
from os import path
import json
from pandas.io.json import json_normalize
import numpy as np

import plotly.graph_objects as go
from plotly.subplots import make_subplots

#python-docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

#Matplotlib
import matplotlib.pyplot as plt
from matplotlib.patches import Shadow

#Color Generator
from faker import Faker
fake = Faker()

#Leyendo parametros
inputFile = str(sys.argv[1])
outputFile = str(sys.argv[2])
host = str(sys.argv[3])
jsonFile = str(sys.argv[4])

#--Loading json file, cleaning useless columns to have an easier json file to work with

with open(jsonFile) as data_file:
    da = json.load(data_file)

for key in da['updates']:

    da['updates'][str(key)]['categories'] = da['updates'][str(key)]['categories'][0]

    if da['updates'][str(key)]['categories'] == "Critical Updates":
        da['updates'][str(key)]['fsorting'] = 1
    elif da['updates'][str(key)] ['categories'] == "Security Updates":
        da['updates'][str(key)]['fsorting'] = 2
    elif da['updates'][str(key)] ['categories'] == "Update Rollups":   
        da['updates'][str(key)]['fsorting'] = 3
    elif da['updates'][str(key)] ['categories'] == "Updates":
        da['updates'][str(key)]['fsorting'] = 4
    elif da['updates'][str(key)] ['categories'] == "Feature Packs":   
        da['updates'][str(key)]['fsorting'] = 5
    elif da['updates'][str(key)] ['categories'] == "Tools":
        da['updates'][str(key)]['fsorting'] = 6
    elif da['updates'][str(key)] ['categories'] == "Application":
        da['updates'][str(key)]['fsorting'] = 7   
    elif da['updates'][str(key)] ['categories'] == "Connectors":
        da['updates'][str(key)]['fsorting'] = 8   
    elif da['updates'][str(key)] ['categories'] == "Definition Updates":
        da['updates'][str(key)]['fsorting'] = 9   
    elif da['updates'][str(key)] ['categories'] == "Developer Kits":
        da['updates'][str(key)]['fsorting'] = 10   
    elif da['updates'][str(key)] ['categories'] == "Guidance":
        da['updates'][str(key)]['fsorting'] = 11

data = json.dumps(da['updates'])
fuckingUpdates = json.loads(data)

for fckupdates in fuckingUpdates.values():
    borrarKB = fckupdates['title']
    kbBorrado = borrarKB.rsplit(' ', 1)[0]
    fckupdates['title'] = kbBorrado

#print(fuckingUpdates) 
lista=[]
for value in fuckingUpdates.values():
    lista.append(value)

#print(lista)

duckingUpdates=sorted(lista, key = lambda i: i['fsorting'])

#TODO Obneter nombre del SO desde AWX
#--Getting Os Name
# iterator = iter(duckingUpdates.values())
# first_value = next(iterator)
# os = first_value['categories'][1]
os = 'Windows Server 2008 R2'

#Verificando archivo a editar
if ( not path.exists(outputFile) ):
    #No existe outputfile, abriendo plantilla
    document = Document(inputFile)
else:
    #Abriendo outputfile
    document = Document(outputFile)

#Si el archivo de salida existe, no se agrega portada
if ( not path.exists(outputFile) ):

    #Creando Portada
    parrafo = document.add_paragraph()
    run = parrafo.add_run('Reporte de Actualizaciones Windows')
    font = run.font
    font.size = Pt(24)
    font.bold = True
    formato_parrafo = parrafo.paragraph_format
    formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    #TODO Obtener de AWX/Ansible el nombre del inventario y colocarlo en la portada
    # parrafo = document.add_paragraph()
    # run = parrafo.add_run('Inventario: Prueba')
    # font = run.font
    # font.size = Pt(24)
    # font.bold = True
    # formato_parrafo = parrafo.paragraph_format
    # formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    parrafo = document.add_paragraph()
    run = parrafo.add_run('Julio 2020')
    font = run.font
    font.size = Pt(24)
    font.bold = True
    formato_parrafo = parrafo.paragraph_format
    formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

#introduciendo informacion

document.add_page_break() #Salto de Pagina

document.add_paragraph() #Enter 

parrafo = document.add_paragraph()
run = parrafo.add_run('Servidor: ' +host)
font = run.font
font.size = Pt(14)
font.bold = True
formato_parrafo = parrafo.paragraph_format
formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph() #Enter

parrafo = document.add_paragraph()
run = parrafo.add_run('Detalles del servidor:')

document.add_paragraph() #Enter

utcmoment_naive = datetime.utcnow()
utcmoment = utcmoment_naive.replace(tzinfo=pytz.utc)
localDatetime = utcmoment.astimezone(pytz.timezone('America/Mexico_City'))

records = (
    ('Operating System', os),
    ('Service Pack', 'None'),
    ('Host Name', host),
    ('Last Status Reported', str(localDatetime.strftime("%d-%m-%Y %H:%M:%S")))
)

table = document.add_table(rows=0, cols=2)
table.style = document.styles['Medium Grid 1 Accent 1']

for id, val in records:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = val

document.add_paragraph() #Enter
document.add_paragraph() #Enter

#Adding Graph iside a table
table_graph = document.add_table(rows=1, cols=1)

#Celda de texto
pic_cells2 = table_graph.rows[0].cells
pic_cells2_pag = pic_cells2[0].add_paragraph()
pic_cells2_pag.alignment=WD_ALIGN_PARAGRAPH.CENTER
pic_cells2_run=pic_cells2_pag.add_run()

document.add_paragraph() #Enter
document.add_paragraph() #Enter

#--Creating table - Autofit didn't work - setting up columns width manually
reportTable = document.add_table(rows=1, cols=4)
#reportTable.autofit = False
#reportTable.style = 'Colorful Shading Accent 6'
reportTable.style = 'Medium Grid 1 Accent 1'
hdr_Cells = reportTable.rows[0].cells
hdr_Cells[0].width = Inches(0.48) 
hdr_Cells[0].text = 'KB'
hdr_Cells[1].width = Inches(4.40)
hdr_Cells[1].text = 'Title'
hdr_Cells[2].width = Inches(0.77)
hdr_Cells[2].text = 'Category'
hdr_Cells[3].width = Inches(0.45)
hdr_Cells[3].text = 'Installed'

#--Filling up table
catgories_dict = {}

for fvalue in duckingUpdates:
    row_Cells = reportTable.add_row().cells
    row_Cells[0].width = Inches(0.48)
    row_Cells[0].text = fvalue['kb']
    row_Cells[1].width = Inches(4.40)
    row_Cells[1].text = fvalue['title']
    row_Cells[2].width = Inches(0.77)
    row_Cells[2].text = fvalue['categories'] #count sobre esta linea
    row_Cells[3].width = Inches(0.45)
    row_Cells[3].text = str(fvalue['installed'])

    if fvalue['categories'] in catgories_dict:
        catgories_dict[fvalue['categories']] = catgories_dict.get(fvalue['categories'])+1
    else:
        catgories_dict[fvalue['categories']] = 1

#Path donde se guardara la grafica como imagen png
#picture_path='/tmp/report_'+host+'2.png'
picture_path='report_'+host+'2.png'

# colores random
fcolor = fake.hex_color()
#Creando vectores para Grafica
titulos = []
contadores = []
explod = []
fcolors = []
for t, c in catgories_dict.items():
    titulos.append(t+' '+str(c))
    contadores.append(c)
    explod.append(0.03)

    if t == 'Critical Updates':
        fcolors.append("#EC1E1E")
    elif t == 'Security Updates':
        fcolors.append("#F3951C")
    elif t == 'Update Rollups':
        fcolors.append("#2672BF")
    elif t == 'Updates':
        fcolors.append("#48B14B")
    elif t == 'Service Packs':
        fcolors.append("#F6F167")
    else:
        fcolors.append(fcolor)

# DONE Modificar grafica con libreria plotly https://plotly.com/python/pie-charts/
# Documentacion: https://plotly.com/python/reference/pie/
#------------------------ INICIO DEL CODIGO DE LA GRAFICA CON PLOTY
labels2=tuple(titulos)
_grafica_ploty = go.Figure(data=[go.Pie(
                             labels=labels2, 
							 values=contadores, 
							 textinfo='percent',
                             insidetextorientation='radial', 
							 pull=explod,
							 marker={"colors":fcolors},
							 textfont={"size":18,"family":"Arial"},
							 direction="counterclockwise",
							 hole=.4,
							 title={
							     "text":"Updates",
							     "font":{"size":24,"family":"Arial"},
							 	 "position":"top center"
							 },
							 outsidetextfont={
							     "size":18,
								 "family":"Arial"
							 },
							 insidetextfont={
							     "size":20,
								 "family":"Arial"
							 },
							 domain={"x":(0,1),"y":(0,1)}
                            )])

_grafica_ploty.update_layout(
    margin=dict(l=0, r=0, t=0, b=0)
)

_grafica_ploty.write_image(picture_path)
#------------------------------- FIN DEL CODIGO DE LA GRAFICA CON MATPLOTLIB

pic_cells2_run.add_picture(picture_path, width=Inches(6))


pic_cells2_run.add_picture(picture_path, width=Inches(4))

#-- Setting font size for the whole table.
for row in reportTable.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        font = run.font
        font.size = Pt(5.5)

document.save(outputFile)