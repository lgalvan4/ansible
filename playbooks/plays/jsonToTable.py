import docx
from docx.shared import Pt
from docx.shared import Inches
import os
import json

with open('report.json') as data_file:
    da = json.load(data_file)
    data = json.dumps(da['updates'])
    duckingUpdates = json.loads(data)

for ftitle in duckingUpdates.values():
    borrarKB = ftitle['title']
    kbBorrado = borrarKB.rsplit(' ', 1)[0]
    ftitle['title'] = kbBorrado

doc = docx.Document()

reportTable = doc.add_table(rows=1, cols=4)
reportTable.autofit = True
reportTable.style = 'Colorful Shading Accent 6'
hdr_Cells = reportTable.rows[0].cells
hdr_Cells[0].width = Inches(0.48)
hdr_Cells[0].text = 'KB'
hdr_Cells[1].width = Inches(4.45)
hdr_Cells[1].text = 'Title'
hdr_Cells[2].width = Inches(0.77)
hdr_Cells[2].text = 'Category'
hdr_Cells[3].width = Inches(0.45)
hdr_Cells[3].text = 'Installed'

for fvalue in duckingUpdates.values():
    row_Cells = reportTable.add_row().cells
    row_Cells[0].width = Inches(0.48)
    row_Cells[0].text = fvalue['kb']
    row_Cells[1].width = Inches(4.45)
    row_Cells[1].text = fvalue['title']
    row_Cells[2].width = Inches(0.77)
    row_Cells[2].text = fvalue['categories'][0]
    row_Cells[3].width = Inches(0.45)
    row_Cells[3].text = str(fvalue['installed'])


for row in reportTable.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        font = run.font
        font.size = Pt(5.5)


doc.save('patchesReport.docx')
